import pdfplumber
from docx import Document
from validate import analyze_text


def read_uploaded_file(uploaded_file):
    print(" fileupload called")

    if uploaded_file is None:
        return ""

    extension = uploaded_file.name.split(".")[-1].lower()

    try:

        # ---------------- TXT ----------------
        if extension == "txt":

            uploaded_file.seek(0)

            try:
                return uploaded_file.read().decode("utf-8")
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                return uploaded_file.read().decode("latin-1")

        # ---------------- PDF ----------------
        elif extension == "pdf":

            uploaded_file.seek(0)

            text = ""

            with pdfplumber.open(uploaded_file) as pdf:
                print("Total Pages:", len(pdf.pages))
                for i, page in enumerate(pdf.pages):

                    page_text = page.extract_text()

                    print(f"Page {i+1}: {repr(page_text)}")

                    if page_text:
                       text += page_text + "\n"
            print("Final PDF Text:",repr(text))
            
            return text

        # ---------------- DOCX ----------------
        elif extension == "docx":

            uploaded_file.seek(0)

            doc = Document(uploaded_file)
            print("Paragraphs:", len(doc.paragraphs))
            text = ""

            for para in doc.paragraphs:
                print("Paragraph:", repr(para.text))
                text += para.text + "\n"

            return text

        else:
            return ""

    except Exception as e:

        print(e)

        return ""



def process_input(user_text, uploaded_file):
    print("process_input called")
    print("=" * 50)
    print("User Text:", repr(user_text))
    print("Uploaded File:", uploaded_file)

    final_text = ""

    if uploaded_file is not None:
        print("Reading file:", uploaded_file.name)

        file_text = read_uploaded_file(uploaded_file)

        print("Extracted text:")
        print(repr(file_text))

        final_text += file_text

    if user_text:
        final_text += "\n" + user_text

    print("Final text:")
    print(repr(final_text))
    print("=" * 50)

    if final_text.strip() == "":
        return {
            "status": False,
            "message": "Please enter text or upload a file."
        }

    result = analyze_text(final_text)
    result["status"] = True
    result["content"] = final_text

    return result